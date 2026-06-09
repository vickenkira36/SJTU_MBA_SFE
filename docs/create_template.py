#!/usr/bin/env python3
"""
Create a pandoc reference-doc (.docx) template conforming to
SJTU Antai MBA thesis format requirements.

Strategy: Start from pandoc's own default reference doc (clean, no legacy numbering),
then override styles to match Antai format spec.
"""

import subprocess
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

# Step 1: Generate pandoc's default reference doc as base
# Use pandoc's built-in default reference doc which has all styles defined
result = subprocess.run(
    ['pandoc', '--print-default-data-file=reference.docx'],
    capture_output=True
)
with open('/tmp/pandoc-base.docx', 'wb') as f:
    f.write(result.stdout)

doc = Document('/tmp/pandoc-base.docx')

# Step 2: Remove ALL numbering definitions to eliminate bullet dots
numbering_part = doc.part.numbering_part
if numbering_part is not None:
    # Clear all abstractNum and num elements
    numbering_el = numbering_part._element
    for child in list(numbering_el):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('abstractNum', 'num'):
            numbering_el.remove(child)

# Step 3: Page setup —— 对齐官方纲领：上3.5 下4.0 左右2.8cm
section = doc.sections[0]
section.top_margin = Cm(3.5)
section.bottom_margin = Cm(4)
section.left_margin = Cm(2.8)
section.right_margin = Cm(2.8)
section.header_distance = Cm(2.5)
section.footer_distance = Cm(3)
section.page_width = Cm(21)
section.page_height = Cm(29.7)

# Step 4: Header —— 左侧固定校名，右侧用 STYLEREF 域自动取当页所在章名
# STYLEREF "Heading 1" 让 Word 在每页页眉自动显示该页所属的一级标题（章名），
# 实现官方模板"页眉随章变化"的效果，而无需按章切分 section。
header = section.header
header.is_linked_to_previous = False
hp = header.paragraphs[0]
hp.clear()

run_left = hp.add_run("上海交通大学MBA学位论文")
run_left.font.name = "宋体"
run_left.font.size = Pt(9)
run_left._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

hp.add_run("\t")

# 右侧 STYLEREF 域：begin / instrText / separate / placeholder / end
def _add_field_run(paragraph, instr, placeholder, font='宋体', size_pt=9):
    """在段落里追加一个 Word 复杂域。"""
    def mk_run():
        r = parse_xml(
            f'<w:r {nsdecls("w")}><w:rPr>'
            f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="{font}"/>'
            f'<w:sz w:val="{int(size_pt*2)}"/><w:szCs w:val="{int(size_pt*2)}"/>'
            f'</w:rPr></w:r>'
        )
        return r
    r1 = mk_run(); r1.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin" w:dirty="true"/>'))
    r2 = mk_run()
    it = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve">{instr}</w:instrText>')
    r2.append(it)
    r3 = mk_run(); r3.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
    r4 = mk_run(); r4.append(parse_xml(f'<w:t {nsdecls("w")} xml:space="preserve">{placeholder}</w:t>'))
    r5 = mk_run(); r5.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
    for r in (r1, r2, r3, r4, r5):
        paragraph._element.append(r)

# 右侧章名不在模板里写死，由 post_process 按章分节后逐节填充静态章名
# （STYLEREF 域在中文 Word 下匹配样式不稳定，改用分节硬写章名的方案）。

pPr = hp._element.get_or_add_pPr()
# 页眉段落显式清零首行缩进（否则继承 Normal 的 2 字符缩进，左侧不靠边）
pPr.append(parse_xml(
    f'<w:ind {nsdecls("w")} w:firstLine="0" w:firstLineChars="0" w:left="0"/>'
))
# Right-aligned tab at text width（左右各 2.8cm）
text_width_twips = int((21 - 2.8 - 2.8) * 567)
tabs = parse_xml(
    f'<w:tabs {nsdecls("w")}>'
    f'  <w:tab w:val="right" w:pos="{text_width_twips}"/>'
    f'</w:tabs>'
)
pPr.append(tabs)

# Bottom border
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="000000"/>'
    f'</w:pBdr>'
)
pPr.append(pBdr)

# Step 4b: Footer —— 居中 PAGE 页码域
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.clear()
fpPr = fp._element.get_or_add_pPr()
fpPr.append(parse_xml(
    f'<w:ind {nsdecls("w")} w:firstLine="0" w:firstLineChars="0" w:left="0"/>'
))
fpPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>'))
_add_field_run(fp, ' PAGE \\* MERGEFORMAT ', '1', font='Times New Roman', size_pt=10.5)

# Step 5: Normal style (body text) —— 对齐官方模板：小四 12pt
style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(12)  # 小四
style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

pf = style_normal.paragraph_format
pf.line_spacing = 1.5
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.first_line_indent = Pt(24)  # 2 Chinese chars @ 12pt

# Step 6: Heading styles — clear any numPr and set fonts

def get_style_by_id(doc, style_id):
    """Get style by style_id (more reliable than by name in pandoc templates)."""
    for s in doc.styles:
        if s.style_id == style_id:
            return s
    raise KeyError(f"No style with id '{style_id}'")

def clean_heading(style, font_size_pt, bold, alignment, east_asia_font,
                  space_before_pt, space_after_pt, outline_level,
                  page_break_before=False):
    """Configure a heading style, removing any numbering."""
    # Font
    style.font.name = 'Times New Roman'
    style.font.size = Pt(font_size_pt)
    style.font.bold = bold
    style.font.color.rgb = None  # auto/black
    style._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia_font)

    # Paragraph format
    pf = style.paragraph_format
    pf.alignment = alignment
    pf.line_spacing = 1.5
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    pf.first_line_indent = Pt(0)
    pf.keep_with_next = True
    pf.page_break_before = page_break_before

    # Remove numPr if exists (kills the bullet dots)
    pPr = style._element.get_or_add_pPr()
    for numPr in pPr.findall(qn('w:numPr')):
        pPr.remove(numPr)

    # Set outline level
    for old in pPr.findall(qn('w:outlineLvl')):
        pPr.remove(old)
    pPr.append(parse_xml(
        f'<w:outlineLvl {nsdecls("w")} w:val="{outline_level}"/>'
    ))

# Heading 1: 三号黑体加粗居中 (16pt)
_h1 = get_style_by_id(doc, 'Heading1')
clean_heading(_h1,
              font_size_pt=16, bold=True,
              alignment=WD_ALIGN_PARAGRAPH.CENTER,
              east_asia_font='黑体',
              space_before_pt=0, space_after_pt=24,
              outline_level=0, page_break_before=True)
# 给 Heading1 补中文别名"标题 1"，让 STYLEREF 按中文名查找也能命中（双保险）
_h1_el = _h1.element
if _h1_el.find(qn('w:aliases')) is None:
    _name_el = _h1_el.find(qn('w:name'))
    _aliases = parse_xml(f'<w:aliases {nsdecls("w")} w:val="标题 1"/>')
    _name_el.addnext(_aliases)

# Heading 2: 四号黑体加粗 (14pt)
clean_heading(get_style_by_id(doc, 'Heading2'),
              font_size_pt=14, bold=True,
              alignment=WD_ALIGN_PARAGRAPH.LEFT,
              east_asia_font='黑体',
              space_before_pt=12, space_after_pt=6,
              outline_level=1)

# Heading 3: 小四号黑体加粗 (12pt)
clean_heading(get_style_by_id(doc, 'Heading3'),
              font_size_pt=12, bold=True,
              alignment=WD_ALIGN_PARAGRAPH.LEFT,
              east_asia_font='黑体',
              space_before_pt=6, space_after_pt=3,
              outline_level=2)

# Step 7: First Paragraph & Body Text — keep indent
for sid, sname in [('FirstParagraph', 'First Paragraph'), ('BodyText', 'Body Text')]:
    try:
        s = get_style_by_id(doc, sid)
    except KeyError:
        try:
            s = doc.styles[sname]
        except KeyError:
            s = doc.styles.add_style(sname, 1)
    s.base_style = style_normal
    s.paragraph_format.first_line_indent = Pt(24)

# Step 8: Also clean heading styles in the generated paragraphs
for p in doc.paragraphs:
    pPr = p._element.find(qn('w:pPr'))
    if pPr is not None:
        for numPr in pPr.findall(qn('w:numPr')):
            pPr.remove(numPr)

# Step 9: Save
output = 'docs/antai-template.docx'
doc.save(output)
os.remove('/tmp/pandoc-base.docx')
print(f"Template saved: {output}")
